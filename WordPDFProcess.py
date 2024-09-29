from tqdm import tqdm
import PyPDF2
import os

def extract_text_from_pdf(pdf_path):
    """
    从指定PDF路径中提取文本内容。
    
    :param pdf_path: PDF文件的路径
    :return: PDF文档的文本内容
    """
    text_content = ''
    try:
        # 打开PDF文件
        with open(pdf_path, 'rb') as file:
            # 创建PdfFileReader对象
            pdf_reader = PyPDF2.PdfReader(file)
            
            # 遍历每一页
            for page_num in range(len(pdf_reader.pages)):
                # 从每一页中提取文本
                page_obj = pdf_reader.pages[page_num]
                text_content += page_obj.extract_text()
    except Exception as e:
        print("Pdf read failed")
    return text_content


def find_files_with_all_strings(root_dir, target_string_list:list[str],exclude_string_list:list[str]):
    """
    检索指定文件夹下的所有文件，并返回文件名包含列表中所有字符串的文件列表。

    要求某些字符串必须有，并排除有某些字符串的文件名
    """
    matching_files = []

    # 遍历给定目录及其子目录中的所有文件和文件夹
    for filename in os.listdir(root_dir):
        # 检查每个文件名是否包含所有目标字符串
        # print(filename)
        if all(string in filename for string in target_string_list) and all(string not in filename for string in exclude_string_list):
            matching_files.append(filename)

    return matching_files


def rename_folders(root_dir, target_string):
    """
    递归查找指定目录及其子目录中名称包含特定字符串的文件夹，并将其重命名，去掉该字符串。

    :param root_dir: 要搜索的根目录路径
    :param target_string: 需要去掉的目标字符串
    """

    def process_directory(current_dir):
        """
        处理当前目录及其子目录中的文件夹。
        """
        # 获取当前目录下的所有文件和文件夹
        for name in os.listdir(current_dir):
            full_path = os.path.join(current_dir, name)
            
            # 如果是文件夹且名称包含目标字符串
            if os.path.isdir(full_path) and target_string in name:
                new_name = name.replace(target_string, '')
                new_full_path = os.path.join(current_dir, new_name)
                
                # 重命名文件夹
                os.rename(full_path, new_full_path)
                print(f'Renamed: {full_path} -> {new_full_path}')
            
            # 递归处理子目录
            if os.path.isdir(full_path):
                process_directory(full_path)

    # 从根目录开始处理
    process_directory(root_dir)
    
    
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def extract_column(doc_path, column_index):
    "提取word文件中第一个表格的某一列"
    # 打开文档
    doc = Document(doc_path)
    
    # 创建一个列表来保存提取的列数据
    extracted_data = []
    
    # 遍历文档中的每个表格
    for table in doc.tables:
        # 对于每个表格，遍历每一行
        for row in table.rows:
            # 获取指定列的单元格
            cell = row.cells[column_index]
            # 将单元格内容添加到列表中
            extracted_data.append(cell.text)
        break  # 实际上只处理第一个表格
    
    return extracted_data

def modify_cell_value(doc_path, table_index, row_index, col_index, new_value):
    "修改一个word当中某一个表格的一个单元格"
    # 打开文档
    doc = Document(doc_path)
    
    # 检查文档中是否有足够的表格
    if table_index >= len(doc.tables):
        raise IndexError(f"文档中没有索引为 {table_index} 的表格")
    
    # 获取指定的表格
    table = doc.tables[table_index]
    
    # 检查表格中是否有足够的行和列
    if row_index >= len(table.rows) or col_index >= len(table.columns):
        raise IndexError(f"表格中没有索引为 ({row_index}, {col_index}) 的单元格")
    
    # 获取指定单元格
    cell = table.cell(row_index, col_index)
    
    # 清空单元格内容
    for paragraph in cell.paragraphs:
        paragraph.clear()
    
    # 如果单元格中已经有段落，直接修改第一个段落的内容
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
        paragraph.text = new_value
    else:
        # 如果单元格中没有段落，添加一个新的段落
        paragraph = cell.add_paragraph(new_value)
    
    # 设置段落对齐方式为居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 保存文档
    doc.save(doc_path)

from docx2pdf import convert

def batch_convert_to_pdf(input_folder, output_folder):
    "批量将一个文件夹下所有的word导出为另一个文件夹下的pdf"
    # 确保输出文件夹存在
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 遍历输入文件夹中的所有文件
    for filename in os.listdir(input_folder):
        input_path = os.path.join(input_folder, filename)
        
        # 仅处理文件，确保是 .docx 格式
        if os.path.isfile(input_path) and filename.endswith(".docx"):
            output_path = os.path.join(output_folder, filename.replace(".docx", ".pdf"))

            # 检查输出文件是否已存在
            if os.path.exists(output_path):
                print(f"Output file {output_path} already exists. Skipping.")
                continue
            
            try:
                # 转换文件
                convert(input_path, output_path)
                print(f"Converted {input_path} to {output_path}")
            except Exception as e:
                print(f"Error converting {input_path}: {e}")